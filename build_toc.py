"""
Examples:
    python build_toc.py --file deposition.pdf --output outputs/toc.json
    python build_toc.py --file inputs/deposition.pdf --output outputs/toc.md
    python build_toc.py --file deposition.pdf --output outputs/toc.docx

json, md, or docx
"""

import argparse
import asyncio
import json
import os
import sys
import time
from pathlib import Path
from dotenv import load_dotenv

from utils.pdf_processor import process_pdf
from utils.metadata_extractor import extract_metadata
from utils.toc_generator import generate_toc

try:
    from docx import Document as DocxDocument 
except ImportError:
    DocxDocument = None

load_dotenv()


def resolve_input(path_str: str) -> Path:
    p = Path(path_str)
    if p.exists():
        return p
    
    alt = Path("inputs") / path_str
    return alt


def ensure_output_dir(path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)


def toc_to_markdown(toc: list) -> str:
    lines = ["# Deposition Table of Contents", ""]
    for i, entry in enumerate(toc, 1):
        topic = entry.get("topic", "Topic")
        page = entry.get("page_start")
        line_start = entry.get("line_start")
        lines.append(f"{i}. **{topic}** - Page {page} - Line {line_start}")
        key_points = entry.get("key_points") or []
        for kp in key_points:
            lines.append(f"   - {kp}")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def toc_to_docx(toc: list, path: Path):
    if DocxDocument is None:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")
    doc = DocxDocument()
    doc.add_heading('Deposition Table of Contents', level=1)
    for i, entry in enumerate(toc, 1):
        topic = entry.get('topic', 'Topic')
        page = entry.get('page_start')
        line_start = entry.get('line_start')
        p = doc.add_paragraph()
        p.add_run(f"{i}. {topic} - Page {page} - Line {line_start}").bold = True
        for kp in entry.get('key_points') or []:
            doc.add_paragraph(kp, style='List Bullet')
    doc.save(path)


async def run_pipeline(input_pdf: Path) -> list:
    print("\nProcessing PDF...")
    nodes = process_pdf(str(input_pdf))
    print(f"Parsed {len(nodes)} chunks")

    if not nodes:
        print("No chunks produced; aborting.")
        return []

    print("\nExtracting metadata...")
    try:
        nodes_with_meta = await extract_metadata(nodes, batch_size=5)
    except Exception as e:
        print(f"Metadata extraction failed ({e}); falling back to basic TOC.")
        nodes_with_meta = []

    print(f"Nodes with deposition_summary: {len(nodes_with_meta)}")
    toc = await generate_toc(nodes_with_meta or nodes)
    print(f"TOC entries: {len(toc)}")
    return toc


def main():
    parser = argparse.ArgumentParser(description="Generate a deposition TOC")
    parser.add_argument("--file", required=True, help="Input deposition PDF path or name in inputs/")
    parser.add_argument("--output", default="outputs/toc.json", help="Output file (.json, .md, or .docx)")
    args = parser.parse_args()

    start = time.time()
    input_path = resolve_input(args.file)
    if not input_path.exists():
        print(f"Error: input file not found: {input_path}")
        sys.exit(1)

    output_path = Path(args.output)
    if output_path.suffix.lower() not in {".json", ".md", ".docx"}:
        print("Error: --output must end with .json, .md, or .docx")
        sys.exit(1)
    ensure_output_dir(output_path)

    print(f"Input:  {input_path}")
    print(f"Output: {output_path}")

    toc = asyncio.run(run_pipeline(input_path))

    ext = output_path.suffix.lower()
    if ext == ".docx":
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(toc, f, indent=2)
    elif ext == ".md":
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(toc_to_markdown(toc))
    else:  # docx
        try:
            toc_to_docx(toc, output_path)
        except RuntimeError as e:
            print(f"DOCX generation error: {e}")
            print("\n Writing to JSON")
            json.dump(toc, f, indent=2)
            sys.exit(1)

    print(f"Written TOC to {output_path}")
    print(f"Completed in {time.time() - start:.2f}s")


if __name__ == "__main__":
    main()
import json
from docx import Document

# json loader
def load_topics(json_file):
    with open(json_file, "r") as f:
        return json.load(f)



# md toc from json
def generate_toc_markdown(topics, output_file):
    with open(output_file, "w") as f:
        f.write("# Table of Contents\n\n")
        for page, page_topics in topics.items():
            f.write(f"## {page}\n")  # header for each page (# in MD)
            for topic in page_topics:
                # apparently markdown doesnt support "·" so i'm just gonna use a hyphen here
                # the first hyphen is for the list item, the rest are to separate the topic, page, and line
                f.write(f"- {topic['topic']} - Page {topic['page_start']} - Line {topic['line_start']}\n")

    print(f"Markdown TOC saved to {output_file}")

# docx toc (word compatible)) from json
def generate_toc_docx(topics, output_file):
    doc = Document()
    doc.add_heading("Table of Contents", level=1)
    for page, page_topics in topics.items():
        doc.add_heading(page, level=2)  # header for each page
        for topic in page_topics:
            # we can use "·" for word (specified in problem statement!)
            doc.add_paragraph(f"{topic['topic']} · Page {topic['page_start']} · Line {topic['line_start']}")
    doc.save(output_file)

    print(f"Word TOC saved to {output_file}")




# main func
if __name__ == "__main__":
    topics_file = "./outputs/extracted_topics.json"
    topics = load_topics(topics_file)

    markdown_output = "./outputs/table_of_contents.md"
    generate_toc_markdown(topics, markdown_output)

    docx_output = "./outputs/table_of_contents.docx"
    generate_toc_docx(topics, docx_output)





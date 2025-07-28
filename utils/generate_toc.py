import json
import sys
from docx import Document

def load_topics(json_file):
    with open(json_file, "r") as f:
        return json.load(f)

def generate_toc_markdown(topics, output_file):
    with open(output_file, "w") as f:
        f.write("# Table of Contents\n\n")
        for page, page_topics in sorted(topics.items(), key=lambda x: int(x[0].split()[1])):
            f.write(f"## {page}\n")
            for topic in page_topics:
                f.write(f"- {topic['topic']} - Page {topic['page_start']} - Line {topic['line_start']}\n")

def generate_toc_docx(topics, output_file):
    doc = Document()
    doc.add_heading("Table of Contents", level=1)
    for page, page_topics in sorted(topics.items(), key=lambda x: int(x[0].split()[1])):
        doc.add_heading(page, level=2)
        for topic in page_topics:
            p = doc.add_paragraph()
            p.add_run(topic['topic']).bold = True
            p.add_run(f" · p{topic['page_start']}, line {topic['line_start']}")
    doc.save(output_file)

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Error: Input JSON file and output file must be specified")
        sys.exit(1)
    
    input_json = sys.argv[1]
    output_file = sys.argv[2]
    
    if not output_file.endswith(('.md', '.docx')):
        print("Error: Output file must end with .md or .docx")
        sys.exit(1)
    
    try:
        topics = load_topics(input_json)
        if output_file.endswith('.md'):
            generate_toc_markdown(topics, output_file)
        else:
            generate_toc_docx(topics, output_file)
        print(f"TOC generated successfully at {output_file}")
    except Exception as e:
        print(f"Error: {str(e)}")
        sys.exit(1)